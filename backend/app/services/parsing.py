"""Local PDF/DOCX parsing (P1).

Docling was spiked and rejected for P1: its default pipeline unconditionally
downloads layout/OCR model weights from HuggingFace Hub on first use, even
with OCR disabled and for a purely text-native PDF (see
IMPLEMENTATION_PHASE_PLAN.md P1 notes). That network dependency has no
pre-provisioning story yet (that lands in P6), so P1 uses `pypdf` and
`python-docx` instead. Docling remains the documented future parser target
once artifact pre-provisioning exists.

Language detection is not implemented in P1 (no detection library was
approved for this phase); `document_language` is fixed to `"en"` here except
for empty extracted text, matching the English-only P1 fixture. Real
detection is deferred to a later phase.
"""

import zipfile
from dataclasses import dataclass

import docx
import pypdf
from docx.opc.exceptions import PackageNotFoundError

PARSER_NAME = "pypdf+python-docx"
PARSER_VERSION = f"pypdf={pypdf.__version__};python-docx={docx.__version__ if hasattr(docx, '__version__') else 'unknown'}"

MAX_PAGE_COUNT = 100
_DOCX_WORDS_PER_PAGE_ESTIMATE = 500
_OLE_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"


class ParsingError(Exception):
    """Base class for typed parsing failures; subclasses map to API error codes."""


class EncryptedDocumentError(ParsingError):
    pass


class DocumentTooLongError(ParsingError):
    pass


class DocumentParseFailedError(ParsingError):
    pass


@dataclass(frozen=True)
class ParsedDocument:
    text: str
    page_count: int
    language: str
    # Character offset in `text` where each page begins (0-indexed list,
    # page N is 1-based = index N-1). PDF offsets are real page boundaries;
    # DOCX offsets are estimated from the same word-count-per-page heuristic
    # used for the page_count estimate (see module docstring). Used only
    # in-memory by the P2 rule engine within the same job-runner pass to
    # attribute clause segments to pages; not persisted.
    page_boundaries: list[int]


def page_for_offset(offset: int, page_boundaries: list[int]) -> int:
    """1-based page number containing the given character offset."""
    page = 1
    for start in page_boundaries[1:]:
        if offset < start:
            break
        page += 1
    return page


def parse_document(path: str, extension: str) -> ParsedDocument:
    if extension == ".pdf":
        return _parse_pdf(path)
    if extension == ".docx":
        return _parse_docx(path)
    raise DocumentParseFailedError(f"Unsupported extension for parsing: {extension}")


def _parse_pdf(path: str) -> ParsedDocument:
    try:
        reader = pypdf.PdfReader(path)
    except Exception as exc:
        raise DocumentParseFailedError(str(exc)) from exc

    if reader.is_encrypted:
        raise EncryptedDocumentError("PDF is encrypted or password-protected.")

    page_count = len(reader.pages)
    if page_count > MAX_PAGE_COUNT:
        raise DocumentTooLongError(f"PDF has {page_count} pages, exceeding the {MAX_PAGE_COUNT}-page limit.")

    try:
        page_texts = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:
        raise DocumentParseFailedError(str(exc)) from exc

    text = "\n".join(page_texts)
    page_boundaries: list[int] = []
    offset = 0
    for page_text in page_texts:
        page_boundaries.append(offset)
        offset += len(page_text) + 1  # +1 for the "\n" joiner

    return ParsedDocument(
        text=text, page_count=page_count, language=_guess_language(text), page_boundaries=page_boundaries
    )


def _parse_docx(path: str) -> ParsedDocument:
    if not zipfile.is_zipfile(path):
        if _looks_encrypted_ole(path):
            raise EncryptedDocumentError("DOCX is encrypted or password-protected.")
        raise DocumentParseFailedError("DOCX file is not a valid OOXML package.")

    try:
        document = docx.Document(path)
    except PackageNotFoundError as exc:
        raise DocumentParseFailedError(str(exc)) from exc
    except Exception as exc:
        raise DocumentParseFailedError(str(exc)) from exc

    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    text = "\n".join(paragraphs)

    # DOCX has no reliable page count without rendering; approximate from
    # word count for the 100-page limit check and for page_boundaries below.
    # This is a documented simplification, not an authoritative page count.
    word_count = len(text.split())
    page_count = max(1, -(-word_count // _DOCX_WORDS_PER_PAGE_ESTIMATE))
    if page_count > MAX_PAGE_COUNT:
        raise DocumentTooLongError(
            f"DOCX estimated at {page_count} pages, exceeding the {MAX_PAGE_COUNT}-page limit."
        )

    page_boundaries = _docx_page_boundaries(paragraphs)

    return ParsedDocument(
        text=text, page_count=page_count, language=_guess_language(text), page_boundaries=page_boundaries
    )


def _docx_page_boundaries(paragraphs: list[str]) -> list[int]:
    boundaries = [0]
    offset = 0
    cumulative_words = 0
    for paragraph in paragraphs:
        page_before = 1 + cumulative_words // _DOCX_WORDS_PER_PAGE_ESTIMATE
        cumulative_words += len(paragraph.split())
        page_after = 1 + cumulative_words // _DOCX_WORDS_PER_PAGE_ESTIMATE
        if page_after > page_before:
            boundaries.append(offset)
        offset += len(paragraph) + 1  # +1 for the "\n" joiner
    return boundaries


def _looks_encrypted_ole(path: str) -> bool:
    try:
        with open(path, "rb") as f:
            header = f.read(8)
    except OSError:
        return False
    return header == _OLE_MAGIC


def _guess_language(text: str) -> str:
    if not text.strip():
        return "unknown"
    return "en"
