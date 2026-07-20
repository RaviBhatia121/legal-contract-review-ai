"""Generate additional synthetic demo contracts for manual testing.

Run from the repository root:
    python scripts/generate_additional_test_contracts.py

Outputs:
- fixtures/compliant-defense-services-agreement.pdf
- fixtures/compliant-defense-services-agreement.docx
- fixtures/nusantara-enterprise-master-services-agreement.pdf
- fixtures/nusantara-enterprise-master-services-agreement.docx

All text is synthetic and intentionally uses the section-heading convention
recognized by the deterministic P2/P3 fallback pipeline.
"""

import os
from collections.abc import Iterable

import docx
from docx.enum.text import WD_BREAK
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas

FIXTURES_DIR = os.path.join(os.path.dirname(__file__), "..", "fixtures")

ContractSection = tuple[str, list[str], bool]


def _wrap_text(c: canvas.Canvas, text: str, font: str, size: int, max_width: float) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if c.stringWidth(candidate, font, size) <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def build_docx(path: str, title: str, sections: Iterable[ContractSection]) -> None:
    document = docx.Document()
    document.add_heading(title, level=0)
    for heading, paragraphs, page_break in sections:
        if page_break:
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        document.add_heading(heading, level=1)
        for para in paragraphs:
            document.add_paragraph(para)
    document.save(path)


def build_pdf(path: str, title: str, sections: Iterable[ContractSection]) -> None:
    c = canvas.Canvas(path, pagesize=LETTER)
    width, height = LETTER
    left_margin = 0.85 * inch
    top_margin = height - 0.8 * inch
    bottom_margin = 0.75 * inch
    line_height = 13
    max_width = width - 2 * left_margin
    y = top_margin

    def new_page() -> None:
        nonlocal y
        c.showPage()
        y = top_margin

    c.setFont("Helvetica-Bold", 16)
    c.drawString(left_margin, y, title)
    y -= 2 * line_height

    for heading, paragraphs, page_break in sections:
        if page_break:
            new_page()
        if y < bottom_margin + 5 * line_height:
            new_page()

        c.setFont("Helvetica-Bold", 12)
        c.drawString(left_margin, y, heading)
        y -= 1.5 * line_height

        c.setFont("Helvetica", 9.5)
        for para in paragraphs:
            for line in _wrap_text(c, para, "Helvetica", 9.5, max_width):
                if y < bottom_margin:
                    new_page()
                    c.setFont("Helvetica", 9.5)
                c.drawString(left_margin, y, line)
                y -= line_height
            y -= line_height * 0.45

    c.save()


COMPLIANT_TITLE = "Compliant Defense Services Support Agreement"

COMPLIANT_SECTIONS: list[ContractSection] = [
    (
        "Parties and Recitals",
        [
            "This Compliant Defense Services Support Agreement is entered into by Garuda Defense Services Authority, a fictional Indonesian defense-services buyer, and Archipelago Secure Systems Ltd., a fictional supplier.",
            "The parties intend this synthetic agreement to represent a clean baseline for managed support, maintenance, and secure data handling services in a regulated defense operating environment.",
        ],
        False,
    ),
    (
        "Section 1. Definitions",
        [
            "Confidential Information means non-public technical, operational, procurement, security, and commercial information disclosed in connection with the services.",
            "Customer Data means all information, records, telemetry, logs, service tickets, and support materials submitted by Buyer or generated from Buyer's systems.",
            "Sensitive Data means Customer Data designated by Buyer as sensitive, restricted, mission-related, export-controlled, personal, or security-relevant.",
        ],
        False,
    ),
    (
        "Section 2. Scope of Services",
        [
            "Supplier shall provide managed support, incident triage, configuration assistance, maintenance coordination, release support, and technical reporting for the platforms listed in Schedule A.",
            "Supplier shall perform the services using trained personnel, documented procedures, controlled tooling, and the service levels agreed by the parties.",
        ],
        False,
    ),
    (
        "Section 6. Data Handling",
        [
            "6.1 Supplier shall process Customer Data only within approved client-controlled systems and only as necessary to perform the services under Buyer's documented instructions.",
            "6.2 Supplier shall not use external or public cloud services for Customer Data unless Buyer gives prior written consent in a signed amendment.",
            "6.3 Supplier shall maintain a current record of processing activities covering location, retention, return, and verified deletion of Customer Data and shall make that record available to Buyer upon request.",
            "6.4 Supplier shall not transfer Customer Data outside the approved operating region without Buyer's prior written consent.",
        ],
        True,
    ),
    (
        "Section 8. Security",
        [
            "8.1 Supplier shall maintain administrative, technical, and physical safeguards consistent with industry standards to protect Customer Data.",
            "8.2 Supplier shall notify Buyer within 24 hours after confirming any security incident affecting Customer Data or service environments.",
            "8.3 Supplier shall maintain documented access control, encryption, logging, and vulnerability management procedures for the services.",
        ],
        True,
    ),
    (
        "Section 9. Confidentiality",
        [
            "9.1 Each party shall protect the other party's Confidential Information using at least reasonable care and no less than the care used for its own comparable information.",
            "9.2 The Supplier may disclose Confidential Information to its Affiliates only on a need-to-know basis and only where such Affiliates are bound by equivalent confidentiality obligations.",
            "9.3 Confidential Information shall not be disclosed publicly or to unrelated third parties except with Buyer's prior written consent or as required by law.",
        ],
        False,
    ),
    (
        "Section 10. Audit and Inspection",
        [
            "10.1 Customer has an audit right to inspect Supplier's records, controls, facilities, and systems relevant to the services upon reasonable notice.",
            "10.2 Supplier shall provide independent evidence, including third-party assessment reports, control test results, and remediation records, when reasonably requested by Customer.",
        ],
        False,
    ),
    (
        "Section 11. Subcontractors",
        [
            "11.1 The Supplier may grant a subcontractor access to Sensitive Data only with Buyer's prior written approval.",
            "11.2 Any subcontractor engaged by Supplier shall be bound by security, confidentiality, audit, and data-protection obligations no less protective than those in this Agreement.",
            "11.3 Supplier remains fully responsible for subcontractor acts and omissions as if they were Supplier's own acts and omissions.",
        ],
        True,
    ),
    (
        "Section 13. Intellectual Property",
        [
            "13.1 The Buyer owns all bespoke Deliverables created under this Agreement. The Supplier retains ownership of pre-existing Background IP and grants the Buyer a perpetual, irrevocable license to use Background IP embedded in the Deliverables.",
        ],
        True,
    ),
    (
        "Section 15. Liability",
        [
            "15.1 Each party's aggregate liability shall not exceed the fees paid in the preceding twelve months.",
            "15.2 The liability cap does not apply to confidentiality breaches, security or data incidents, intellectual property infringement, fraud, or wilful misconduct.",
            "15.3 Supplier shall indemnify Buyer against third-party claims that a Deliverable infringes a third party's intellectual property rights.",
        ],
        True,
    ),
    (
        "Section 17. Termination",
        [
            "17.1 Either party may terminate this Agreement for uncured material breach upon thirty days' written notice.",
            "17.2 Upon termination for material breach or a security incident, the Supplier shall return or verifiably delete all Customer Data and provide reasonable transition assistance.",
        ],
        True,
    ),
    (
        "Section 18. General",
        [
            "This synthetic agreement is governed by the law selected in Schedule C. Notices must be in writing. Amendments require signed written agreement by authorized representatives of both parties.",
        ],
        False,
    ),
]


ENTERPRISE_TITLE = "Nusantara Enterprise Master Services Agreement"

ENTERPRISE_SECTIONS: list[ContractSection] = [
    (
        "Parties and Recitals",
        [
            "This Nusantara Enterprise Master Services Agreement is entered into by Nusantara Integrated Defense Services Authority, a fictional Indonesian defense-sector services organization, and Maritime Systems Operations Pte. Ltd., a fictional enterprise technology supplier.",
            "The agreement establishes a master framework for multiple statements of work covering platform support, integration services, service desk operations, field maintenance coordination, reporting, and operational analytics for non-combat defense logistics environments.",
            "The parties acknowledge that the services may involve sensitive operational metadata, protected procurement records, configuration data, incident information, and support materials that require controlled handling throughout the service lifecycle.",
        ],
        False,
    ),
    (
        "Section 1. Definitions",
        [
            "Customer Data means all data, records, documents, logs, configuration files, support tickets, service reports, access records, diagnostic extracts, and derivative records submitted by or on behalf of Buyer.",
            "Sensitive Data means Customer Data categorized by Buyer as restricted, mission-sensitive, regulated, export-controlled, personal, security-relevant, or otherwise requiring enhanced handling controls.",
            "Deliverables means work product, configuration, documentation, scripts, reports, interfaces, dashboards, and other materials created specifically for Buyer under a statement of work.",
        ],
        False,
    ),
    (
        "Section 2. Master Services Framework",
        [
            "Supplier shall provide services only under executed statements of work. Each statement of work shall identify scope, milestones, deliverables, service levels, acceptance criteria, dependencies, environments, reporting obligations, and authorized Buyer points of contact.",
            "Supplier shall maintain a governance plan covering onboarding, role assignment, issue escalation, change control, risk reporting, dependency management, and monthly service review meetings with Buyer.",
            "If a statement of work conflicts with this Agreement, this Agreement controls unless the statement of work expressly identifies the specific clause being modified and is signed by both parties' authorized representatives.",
        ],
        False,
    ),
    (
        "Section 3. Service Management",
        [
            "Supplier shall operate a support desk, maintain incident queues, track service requests, provide root-cause reports for major incidents, and maintain a corrective-action register for repeat issues.",
            "Supplier shall meet the service levels in Schedule B, including response times, resolution targets, reporting cadence, and service-credit calculations for chronic missed commitments.",
            "Supplier shall not change production configurations, network rules, privileged access, or integration endpoints except through the approved change-control process.",
        ],
        True,
    ),
    (
        "Section 4. Governance and Compliance",
        [
            "Supplier shall comply with applicable procurement, anti-bribery, sanctions, export-control, personal-data, cybersecurity, and records-retention requirements that apply to the services.",
            "Supplier shall maintain training records for personnel assigned to Buyer work and shall remove any personnel who fail background, confidentiality, or security requirements reasonably imposed by Buyer.",
            "Supplier shall provide monthly operational reports covering open incidents, change requests, access reviews, subcontractor participation, control exceptions, and remediation status.",
        ],
        False,
    ),
    (
        "Section 6. Data Handling",
        [
            "6.1 Supplier shall process Customer Data only within approved client-controlled systems and only as necessary to perform the services under Buyer's documented instructions.",
            "6.2 Supplier shall not use external or public cloud services for Customer Data unless Buyer gives prior written consent in a signed amendment.",
            "6.3 Supplier shall maintain a current record of processing activities covering location, retention, return, and verified deletion of Customer Data and shall make that record available to Buyer upon request.",
            "6.4 Supplier shall not transfer Customer Data outside the approved operating region without Buyer's prior written consent.",
        ],
        True,
    ),
    (
        "Section 7. Business Continuity",
        [
            "Supplier shall maintain a business continuity and disaster recovery plan proportionate to the services and shall test the plan at least annually.",
            "Supplier shall notify Buyer of material continuity events, failed recovery tests, or unresolved continuity gaps that could affect service availability or data protection.",
            "Supplier shall coordinate continuity exercises with Buyer when a service supports critical operational reporting or time-sensitive logistics processes.",
        ],
        False,
    ),
    (
        "Section 8. Security",
        [
            "8.1 Supplier shall maintain administrative, technical, and physical safeguards consistent with industry standards to protect Customer Data.",
            "8.2 Supplier shall notify Buyer within 48 hours after confirming any security incident affecting Customer Data or service environments.",
            "8.3 Supplier shall maintain documented access control, encryption, logging, and vulnerability management procedures for the services.",
            "8.4 Supplier shall review privileged access quarterly, revoke access promptly when no longer required, and retain security logs for investigation and audit purposes.",
        ],
        True,
    ),
    (
        "Section 9. Confidentiality",
        [
            "9.1 Each party shall protect the other party's Confidential Information using at least reasonable care and no less than the care used for its own comparable information.",
            "9.2 The Supplier may disclose Confidential Information to its Affiliates only on a need-to-know basis and only where such Affiliates are bound by equivalent confidentiality obligations.",
            "9.3 Confidential Information includes operational records, technical architecture, pricing, incident reports, security findings, procurement materials, and any information marked or reasonably understood to be confidential.",
        ],
        False,
    ),
    (
        "Section 10. Audit and Inspection",
        [
            "10.1 Supplier shall provide vendor-selected summary only reports for any audit request and may decline access to underlying control materials.",
            "10.2 Buyer may review the summary reports during quarterly governance meetings, but Supplier is not required to provide third-party assessment records unless Supplier agrees separately.",
        ],
        True,
    ),
    (
        "Section 11. Subcontractors",
        [
            "11.1 The Supplier may grant a subcontractor access to Sensitive Data upon prior notice to the Buyer.",
            "11.2 Any subcontractor engaged by Supplier shall be bound by security and confidentiality obligations no less protective than those in this Agreement.",
            "11.3 Supplier shall maintain a subcontractor register identifying work performed, locations, access categories, and service dependencies, and shall provide the register to Buyer during quarterly governance meetings.",
        ],
        False,
    ),
    (
        "Section 12. Change Control",
        [
            "Each party may request changes through a written change request. A change request must identify business justification, impacted services, cost, timeline, security impact, testing evidence, rollback approach, and approval requirements.",
            "No change is effective unless approved by both parties in writing. Emergency changes may be implemented only where necessary to prevent service disruption or security harm and must be retrospectively documented.",
        ],
        False,
    ),
    (
        "Section 13. Intellectual Property",
        [
            "13.1 The Buyer owns all bespoke Deliverables created under this Agreement. The Supplier retains ownership of pre-existing Background IP and grants the Buyer a perpetual, irrevocable license to use Background IP embedded in the Deliverables.",
            "13.2 Supplier shall not incorporate open-source components, third-party tools, or reusable libraries into Deliverables in a manner that restricts Buyer's internal enterprise use without Buyer's prior written approval.",
        ],
        True,
    ),
    (
        "Section 14. Warranties",
        [
            "Supplier warrants that the services will be performed with reasonable skill and care by qualified personnel and in accordance with the applicable statement of work.",
            "Supplier warrants that Deliverables will materially conform to agreed acceptance criteria for ninety days after acceptance, excluding defects caused by unauthorized Buyer modifications.",
        ],
        False,
    ),
    (
        "Section 15. Liability",
        [
            "15.1 Each party's aggregate liability shall not exceed the fees paid in the preceding twelve months.",
            "15.2 The liability cap does not apply to confidentiality breaches, security or data incidents, intellectual property infringement, fraud, or wilful misconduct.",
            "15.3 Supplier shall indemnify Buyer against third-party claims that a Deliverable infringes a third party's intellectual property rights.",
            "15.4 Neither party excludes liability to the extent such exclusion is prohibited by applicable law.",
        ],
        True,
    ),
    (
        "Section 16. Insurance",
        [
            "Supplier shall maintain commercially reasonable professional liability, cyber liability, employer liability, and general liability insurance during the term and for any required tail period.",
            "Supplier shall provide certificates of insurance upon request and notify Buyer if a policy is cancelled, materially reduced, or not renewed.",
        ],
        False,
    ),
    (
        "Section 17. Termination",
        [
            "17.1 Either party may terminate this Agreement for uncured material breach upon thirty days' written notice.",
            "17.2 Upon termination for material breach or a security incident, the Supplier shall return or verifiably delete all Customer Data and provide reasonable transition assistance.",
            "17.3 Buyer may terminate a statement of work for convenience on sixty days' written notice, subject to payment for accepted services performed before termination.",
        ],
        True,
    ),
    (
        "Section 18. Transition and Exit",
        [
            "Supplier shall cooperate with Buyer and any replacement supplier to support orderly transition, knowledge transfer, data export, access revocation, open-ticket handover, and return of Buyer-controlled materials.",
            "Transition assistance shall be provided at agreed rates unless termination arises from Supplier's uncured material breach or security incident, in which case transition assistance is included for a reasonable period.",
        ],
        False,
    ),
    (
        "Section 19. General",
        [
            "This Agreement, the schedules, and executed statements of work constitute the entire agreement between the parties. Amendments must be in writing and signed by authorized representatives.",
            "Notices must be delivered to the addresses in Schedule D. Neither party may assign this Agreement without prior written consent except to a successor in connection with a permitted corporate reorganization.",
        ],
        False,
    ),
]


def main() -> None:
    os.makedirs(FIXTURES_DIR, exist_ok=True)
    outputs = [
        ("compliant-defense-services-agreement", COMPLIANT_TITLE, COMPLIANT_SECTIONS),
        ("nusantara-enterprise-master-services-agreement", ENTERPRISE_TITLE, ENTERPRISE_SECTIONS),
    ]
    for basename, title, sections in outputs:
        pdf_path = os.path.join(FIXTURES_DIR, f"{basename}.pdf")
        docx_path = os.path.join(FIXTURES_DIR, f"{basename}.docx")
        build_pdf(pdf_path, title, sections)
        build_docx(docx_path, title, sections)
        print(f"Wrote {pdf_path}")
        print(f"Wrote {docx_path}")


if __name__ == "__main__":
    main()
